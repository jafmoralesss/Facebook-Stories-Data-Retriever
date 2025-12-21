from docx import Document
import os

class WordReportView:
    def __init__(self, filename="Reporte_Completo_Facebook.docx"):
        self.filename = filename

    def generate_report(self, data_list):
        if not data_list:
            print("--- Lista de datos vacía. No se generará el reporte. ---")
            return
        
        doc = Document()
        doc.add_heading('Extracción de Estados de Facebook', 0)

        total = len(data_list)
        print(f"--- Escribiendo {total} estados en el documento... ---")
        
        #Bucle para escribir cada estado
        for i, texto in enumerate(data_list, 1):
            doc.add_heading(f'Estado #{i}', level=2)
            doc.add_paragraph(texto)
            doc.add_paragraph("_"* 40)
        
        #Guardar archivo
        doc.save(self.filename)

        #Mostrar ruta absoluta
        ruta_completa = os.path.abspath(self.filename)
        print(f"--- Archivo guardado en: \n{ruta_completa} ---")