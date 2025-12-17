import asyncio
from playwright.async_api import async_playwright
from bs4 import BeautifulSoup
from docx import Document

async def run():
    print("Testing Validation")
    async with async_playwright() as p:
#abriendo navegador
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
#ingreso a web de prueba
        await page.goto("https://example.com")
        html = await page.content()
#procesamiento del HTML
        soup = BeautifulSoup(html, "html.parser")
        texto = soup.find("h1").get_text()
        print(f"Texto extraido: {texto}")

#creando archivo word
        doc=Document()
        doc.add_heading('Resultado de Prueba', 0)
        doc.add_paragraph(f'Se extrajo el texto: {texto}')
        doc.save('prueba_exitosa.docx')
        print ("Archivo 'prueba_exitosa.docx' generado")

        await browser.close()

if __name__ == "__main__":
    asyncio.run(run())