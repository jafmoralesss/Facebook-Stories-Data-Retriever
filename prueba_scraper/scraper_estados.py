import asyncio
import random
from playwright.async_api import async_playwright
from bs4 import BeautifulSoup
from docx import Document

# --- CONFIGURACIÓN ---
URL_OBJETIVO = "https://www.facebook.com/me" 
CANTIDAD_SCROLLS = 3

async def scrapear_facebook():
    print("🤖 Iniciando el Bot (Modo Preciso)...")
    
    async with async_playwright() as p:
        # 1. Navegador y Sesión
        browser = await p.chromium.launch(headless=False)
        context = await browser.new_context(storage_state="facebook_auth.json")
        page = await context.new_page()

        # 2. Ir al perfil
        print(f"🌍 Yéndonos a: {URL_OBJETIVO}")
        await page.goto(URL_OBJETIVO)
        await page.wait_for_selector("div[role='main']", timeout=10000)

        # 3. Scroll (Para cargar más estados)
        print("⬇️  Bajando para cargar publicaciones...")
        for i in range(CANTIDAD_SCROLLS):
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(3000) # Espera 3 segundos entre scrolls

        # 4. Extracción
        print("👀 Leyendo el HTML...")
        content = await page.content()
        soup = BeautifulSoup(content, 'html.parser')

        # --- AQUI ESTA EL CAMBIO IMPORTANTE ---
        # Usamos un selector CSS. Los puntos (.) representan las clases.
        # Esto busca: <div class="x78zum5 xdt5ytf x1rife3k">
        bloques_texto = soup.select('div.x78zum5.xdt5ytf.x1rife3k')
        
        print(f"✨ Se encontraron {len(bloques_texto)} bloques de texto exactos.")

        # 5. Guardar en Word
        doc = Document()
        doc.add_heading('Estados Limpios de Facebook', 0)

        contador = 1
        for bloque in bloques_texto:
            # get_text con separator=" " une líneas con un espacio para que sea legible
            texto_limpio = bloque.get_text(separator="\n", strip=True)

            # Filtro simple: Si tiene menos de 2 letras, probablemente es un emoji suelto o error
            if len(texto_limpio) > 1:
                doc.add_heading(f'Estado #{contador}', level=2)
                doc.add_paragraph(texto_limpio)
                doc.add_paragraph("_" * 30)
                contador += 1

        doc.save('estados_limpios.docx')
        print("✅ Archivo 'estados_limpios.docx' creado con éxito.")

        await browser.close()

if __name__ == "__main__":
    asyncio.run(scrapear_facebook())